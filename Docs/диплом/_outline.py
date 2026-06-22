# -*- coding: utf-8 -*-
import sys
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
for fn in ["Диплом.docx","диплом_ксюша.docx"]:
    d = Document(fn)
    paras = d.paragraphs
    print("\n==================== %s ====================" % fn)
    print("paragraphs: %d, tables: %d" % (len(paras), len(d.tables)))
    # words estimate
    words = sum(len(p.text.split()) for p in paras)
    print("approx words: %d (~%d pages @300w)" % (words, words//300))
    print("---- headings (bold/centered or numbered) ----")
    for p in paras:
        t = p.text.strip()
        if not t: continue
        bold = any(r.bold for r in p.runs if r.text.strip())
        centered = p.paragraph_format.alignment == AL.CENTER
        is_head = (bold and (centered or len(t) < 80)) or t[:1].isdigit() and "." in t[:4] and len(t)<90
        if is_head and len(t) < 110:
            print(("  C " if centered else "    ") + t[:100])
