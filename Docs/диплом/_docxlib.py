# -*- coding: utf-8 -*-
"""Хелперы для наполнения Диплом.docx с оформлением по нормоконтролю АлтГУ
(Times New Roman 14, интервал 1.5, отступ 1.25 см; подписи 12 пт)."""
from docx.shared import Pt, Emu, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Times New Roman"
FLI = Emu(450215)   # абзацный отступ 1.25 см
SB6 = Pt(6)


def _font(r, size=14, bold=False, italic=False):
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    r.font.name = FONT
    rpr = r._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts'); rpr.append(rfonts)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rfonts.set(qn(a), FONT)


def body(doc, text, size=14, justify=True):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = AL.JUSTIFY if justify else AL.LEFT
    pf.line_spacing = 1.5
    pf.first_line_indent = FLI
    pf.space_after = Pt(0)
    _font(p.add_run(text), size=size)
    return p


def body_rich(doc, parts, size=14):
    """parts: список (текст, bold, italic)."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = AL.JUSTIFY
    pf.line_spacing = 1.5
    pf.first_line_indent = FLI
    pf.space_after = Pt(0)
    for tup in parts:
        text = tup[0]; b = tup[1] if len(tup) > 1 else False; it = tup[2] if len(tup) > 2 else False
        _font(p.add_run(text), size=size, bold=b, italic=it)
    return p


def bullet(doc, text, size=14):
    return body(doc, "– " + text, size=size)


def chapter(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = AL.CENTER
    pf.line_spacing = 1.5
    pf.page_break_before = True
    pf.space_after = Pt(0)
    _font(p.add_run(text), size=14, bold=True)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = AL.LEFT
    pf.line_spacing = 1.5
    pf.first_line_indent = FLI
    pf.space_before = SB6
    pf.space_after = Pt(0)
    _font(p.add_run(text), size=14, bold=True)
    return p


def figure(doc, path, caption, width_in=6.2):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = AL.CENTER
    pf.line_spacing = 1.0
    pf.space_before = SB6
    p.add_run().add_picture(path, width=Inches(width_in))
    c = doc.add_paragraph()
    cf = c.paragraph_format
    cf.alignment = AL.CENTER
    cf.line_spacing = 1.0
    cf.space_after = SB6
    _font(c.add_run(caption), size=12)
    return p


def formula(doc, text, number):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = AL.CENTER
    pf.line_spacing = 1.5
    _font(p.add_run(text + "\t\t(" + number + ")"), size=14)
    return p


def table_caption(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = AL.RIGHT
    pf.line_spacing = 1.0
    pf.space_before = SB6
    _font(p.add_run(text), size=12)
    return p


def _mono(r, size=10.5):
    r.font.size = Pt(size)
    r.font.name = "Courier New"
    rpr = r._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts'); rpr.append(rfonts)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rfonts.set(qn(a), "Courier New")


def listing_caption(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = AL.LEFT
    pf.line_spacing = 1.0
    pf.space_before = SB6
    pf.first_line_indent = FLI
    _font(p.add_run(text), size=12)
    return p


def code(doc, text, size=10.5):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.0
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.left_indent = Emu(228600)  # ~0.25 cm
    lines = text.split("\n")
    for i, line in enumerate(lines):
        r = p.add_run(line)
        _mono(r, size)
        if i < len(lines) - 1:
            r.add_break()
    return p


def _borders(table):
    tblPr = table._tbl.tblPr
    b = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), '4')
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), '000000')
        b.append(e)
    tblPr.append(b)


def _shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear'); sh.set(qn('w:color'), 'auto'); sh.set(qn('w:fill'), fill)
    tcPr.append(sh)


def table(doc, rows, header=True, aligns=None, size=12, widths=None):
    nr, nc = len(rows), len(rows[0])
    t = doc.add_table(rows=nr, cols=nc)
    t.alignment = 1  # center
    _borders(t)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.cell(i, j)
            p = cell.paragraphs[0]
            pf = p.paragraph_format
            pf.line_spacing = 1.0
            pf.space_after = Pt(0)
            if header and i == 0:
                pf.alignment = AL.CENTER
            else:
                pf.alignment = (aligns[j] if aligns else AL.LEFT)
            _font(p.add_run(str(val)), size=size, bold=(header and i == 0))
            if header and i == 0:
                _shade(cell, "1e4b8f")
                p.runs[0].font.color.rgb = __import__('docx').shared.RGBColor(0xFF, 0xFF, 0xFF)
            if widths:
                cell.width = Inches(widths[j])
    return t
